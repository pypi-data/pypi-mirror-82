# -*- coding: utf-8 -*-
"""
Created on Mon Jan 30 09:51:02 2017

@author: audun
"""

from pylatex import Document, PageStyle, Head, Foot, MiniPage, \
    StandAloneGraphic, MultiColumn, Tabu, LongTabu, LargeText, MediumText, \
    LineBreak, NewPage, Tabularx, TextColor, simple_page_number, Command, \
    Figure, Package, SubFigure
from pylatex.utils import bold, NoEscape
from pylatex.base_classes import CommandBase, Arguments
import os
from pylatex.utils import make_temp_dir
import uuid
import posixpath
from pylatex import Section, Subsection, Subsubsection, Package
import pandas as pd
import numpy as np
from pyreportlib import get_document
from pyreportlib.utils import excel_to_latex
from docx import Document
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn
import datetime


class SidewaysFigure(Figure):
    packages = [Package('rotating')]


# setting the levels of each section (maximum recursion level: 3)
def _set_section_levels(content):
    """
    _set_section_levels adds a keyword 'level' at each sub_dict with value equal to the recursion
    level of that sub_dict in the dictionary content. The level value is used later to make sure that the content is
    appended at the right level in the report (main chapter, sub chapter and sub sub chapter, etc.)
    """
    for item in content:
        item['level'] = 1
        if item.get('content'):
            for sub_item in item['content']:
                if sub_item.get('content'):
                    sub_item['level'] = 2
                    for sub_sub_item in sub_item['content']:
                        if isinstance(sub_sub_item,dict):
                            if sub_sub_item.get('content'):
                                sub_sub_item['level'] = 3
                                for sub_sub_sub_item in sub_sub_item['content']:
                                    if isinstance(sub_sub_sub_item, dict):
                                        if sub_sub_sub_item.get('content'):
                                            sub_sub_sub_item['level'] = np.nan
                                        else:
                                            sub_sub_sub_item['level'] = 3
                            else:
                                sub_sub_item['level'] = 2
                else:
                    sub_item['level'] = 1
    return content


def _format_content(content, content_key):
    list_content = []
    not_list_content = []
    if isinstance(content[content_key],dict):
        for key, val in content[content_key].items():
            if isinstance(val, list):
                list_content.append(key)
            else:
                not_list_content.append(key)

        if len(list_content) > 0:
            new_content = []
            for i in range(len(content[content_key][list_content[0]])):  # assume that all lists have same lengths
                new_dict = {}
                for key in list_content:
                    new_dict[key] = content[content_key][key][i]
                for key in not_list_content:
                    new_dict[key] = content[content_key][key]
                new_content.append(new_dict)
            content[content_key] = new_content
        else:
            content[content_key] = [content[content_key]]


def _format_document_dict(content):
    """
        _format_document_dict makes sure that all lists in the content dictionary are lists of dicts. This is done to
        make the content specification more flexible for the user.

        #Example
        Say the user sends in five spreadsheets to a chapter.
        Since they all have he same kwargs, it is convenient for the user to only give the kwargs argument once.
        _format_document_dict will return a list of dicts, one for each spreadsheet with the specified kwargs.

        INPUT:
        {
            "table": {
                "filename": [
                    "files/spreadsheet0.xlsx",
                    "files/spreadsheet1.xlsx",
                    "files/spreadsheet2.xlsx",
                    "files/spreadsheet3.xlsx",
                    "files/spreadsheet4.xlsx"
                ],
                "kwargs": {
                    "sheet_name": "table_2",
                    "index_col": 0
                }
            }
        }
        RETURNS:
        {
            "table": [
                {
                    "filename": "files/spreadsheet0.xlsx",
                    "kwargs": {
                        "sheet_name": "table_2", "index_col": 0}},
                {
                    "filename": "files/spreadsheet1.xlsx",
                    "kwargs": {
                        "sheet_name": "table_2", "index_col": 0}},
                {
                   "filename": "files/spreadsheet2.xlsx",
                    "kwargs": {
                        "sheet_name": "table_2", "index_col": 0}},
                {
                    "filename": "files/spreadsheet3.xlsx",
                    "kwargs": {
                        "sheet_name": "table_2", "index_col": 0}},
                {
                    "filename": "files/spreadsheet4.xlsx",
                    "kwargs": {
                        "sheet_name": "table_2", "index_col": 0}}
            ]
        }
    """
    if isinstance(content, list):
        for item in content:
            if item.get('title'):
                _format_document_dict(item['content'])
            else:
                _format_document_dict(item)
    else:
        if content.get('table'):
            _format_content(content, 'table')
        if content.get('image'):
            _format_content(content, 'image')
        if content.get('subimage'):
            _format_content(content, 'subimage')
    return content


def _get_section(**content):
    if content['level'] == 1:
        return Section(content['title'])
    elif content['level'] == 2:
        return Subsection(content['title'])
    elif content['level'] == 3:
        return Subsubsection(content['title'])


def _get_last_section(doc):
    for stuff in doc.data[::-1]:
        if isinstance(stuff, (Section, Subsection, Subsubsection)):
            return stuff


def _add_toc_to_docx(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p


def _append2latexdoc(doc, content):
    if isinstance(content, list):
        for item in content:
            if item.get('title'):
                doc.append(_get_section(**item))
                _append2latexdoc(doc, item['content'])
            else:
                _append2latexdoc(doc, item)
    else:
        section = _get_last_section(doc)
        if content.get('text'):
            if isinstance(content['text'], dict):
                section.append(open(content['text']['filename']).read())
            else:
                temp = ''
                if isinstance(content['text'],list):
                    for subtext in content['text']:
                        temp += subtext
                elif isinstance(content['text'],str):
                    temp = content['text']
                else:
                    raise Exception(f'Did not understand format of text string: \n {content["text"]}')
                section.append(temp)
        if content.get('latex_code'):
            if isinstance(content['latex_code'], dict):
                section.append(NoEscape(open(content['latex_code']['filename']).read()))
            else:
                section.append(NoEscape(content['latex_code']))
        if content.get('table'):
            for table in content['table']:
                section.append(NoEscape('\\begin{table}[H]'))  # note require float latex package for H command
                if table.get('filename'):
                    df = pd.read_excel(table['filename'], **table['kwargs'])
                elif isinstance(table.get('dataframe'),pd.DataFrame):
                    df = table.get('dataframe')
                section.append(NoEscape(df.to_latex(longtable=True,multicolumn_format='c')))
                section.append(NoEscape('\\end{table}'))
        if content.get('image'):
            for image in content.get('image'):
                section.append(NoEscape('\\begin{figure}[H]'))  # note require float latex package for H command
                Figure.add_image(section, image['filename'])
                section.append(NoEscape('\\end{figure}'))
        if content.get('subimage'):
            figure = Figure(position='H')
            for i, subimage in enumerate(content['subimage']):
                subfigure = SubFigure(width=NoEscape(
                    r'{}\linewidth'.format(np.round(1. / subimage.get('nr_horizontal_subimages', 2), 2) - 0.01)))
                subfigure.add_image(subimage['filename'])
                if subimage.get('caption', False):
                    subfigure.add_caption(subimage['caption'])
                if subimage.get('figure_caption', False) and i == 0:
                    figure.add_caption(subimage['figure_caption'])
                figure.append(subfigure)
                if (i + 1) % subimage.get('nr_horizontal_subimages', 2) == 0 and i != 0 or subimage.get(
                        'nr_horizontal_subimages', 2) == 1:
                    section.append(figure)
                    figure = Figure(arguments=NoEscape('\ContinuedFloat'), position='H')
            section.append(figure)
        if content.get('packages'):
            [doc.packages.append(Package(package)) for package in content['packages']]


def _append2worddoc(doc, content):
    if isinstance(content, list):
        for item in content:
            if item.get('title'):
                if item.get('level') == 1: # Add page break for level1-headings
                    doc.add_page_break()
                doc.add_heading(item.get('title'), level=item.get('level'))
                _append2worddoc(doc, item.get('content'))
            else:
                _append2worddoc(doc, item)
    else:
        if content.get('text'):
            if isinstance(content['text'], dict):
                doc.add_paragraph(open(content['text']['filename']).read())
            else:
                doc.add_paragraph(content['text'])
        if content.get('latex_code'):
            print('Latex code not directly supported in word, ignored. ')
        if content.get('table'):
            for table in content['table']:
                if table.get('filename'):
                    df = pd.read_excel(table['filename'], **table['kwargs'])
                elif isinstance(table.get('dataframe'),pd.DataFrame):
                    df = table.get('dataframe')

                # add a table to the end and create a reference variable
                # extra row is so we can add the header row
                t = doc.add_table(df.shape[0] + 1, df.shape[1])

                # add the header rows.
                for j in range(df.shape[-1]):
                    t.cell(0, j).text = df.columns[j]

                # add the rest of the data frame
                for i in range(df.shape[0]):
                    for j in range(df.shape[-1]):
                        t.cell(i + 1, j).text = str(df.round(decimals=5).values[i, j])
                p = doc.add_paragraph()
                r = p.add_run()
                r.add_break()

        if content.get('image'):
            for image in content.get('image'):
                doc.add_picture(image.get('filename'), width=Cm(12))
        if content.get('subimage'):
            print('The subfigure feature is not yet supported by the word compilator, figure is ignored')


def make_latex_document(document_title='Document title', document_filename='default_filename', content=[],
                  **doc_template_kwargs):
    if doc_template_kwargs.get('workflow_ID'):
        wf_id = doc_template_kwargs.get('workflow_ID')

    doc = get_document(document_title, **doc_template_kwargs)
    content = _set_section_levels(content)
    content = _format_document_dict(content)
    _append2latexdoc(doc, content)
    doc.generate_tex(document_filename)
    doc.generate_pdf(document_filename, clean_tex=False, clean=False, compiler='pdfLaTeX')
    doc.generate_pdf(document_filename, clean_tex=False, clean=True, compiler='pdfLaTeX')
    return doc


def make_word_document(document_title='Document title', document_filename='default_filename', content=[],
                  **doc_template_kwargs):

    if doc_template_kwargs.get('document_template'):
        doc = Document(doc_template_kwargs.get('document_template'))
    else:
        doc = Document()

    doc.core_properties.title = document_title

    if doc_template_kwargs.get('workflow_ID'):
        wf_id = doc_template_kwargs.get('workflow_ID')
        for paragraph in doc.paragraphs:
            if 'Disclaimer' in paragraph.text:
                distext = paragraph.text
                distext = distext.replace('xxxxxxxx', f'{wf_id}')
                distext = distext.replace('xx.xx.xxxx', datetime.datetime.now().strftime("%m/%d/%Y, %H:%M:%S"))
                paragraph.text = distext

    content = _set_section_levels(content)
    content = _format_document_dict(content)
    _append2worddoc(doc, content)
    doc.save(f'{document_filename}.docx')
    return doc
    # return content # For development
