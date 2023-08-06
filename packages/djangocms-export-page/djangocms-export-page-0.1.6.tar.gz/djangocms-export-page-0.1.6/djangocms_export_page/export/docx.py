from collections import namedtuple
from io import BytesIO

from docx import Document
from docx.shared import Cm, Pt

from .common import PageExport

InlineStyle = namedtuple('InlineStyle', ['font', 'size', 'bold', 'italic'])


class DocxPageExport(PageExport):
    margin = Cm(2)
    font_name = 'Tahoma'
    table_style = 'TableGrid'
    column_layout = [Cm(2.5), Cm(7), Cm(8)]
    row_sytle = [InlineStyle(None, Pt(8), bold=True, italic=False), None, None]

    def setup_document(self):
        doc = Document()

        doc.styles['Normal'].font.name = self.font_name

        section = doc.sections[0]
        section.left_margin = section.right_margin = self.margin

        return doc

    def export(self):
        self.document = self.setup_document()
        self.document.add_heading(self.object.get_title(), level=2)
        self.document.add_paragraph().add_run(self.page_url).italic = True

        sections = self.get_data()

        for section_name, section, components in sections:
            self.document.add_heading(section_name, level=5)

            for component_name, instance, fields in components:
                rows = [self.get_row(field) for field in fields]

                if rows:
                    self.document.add_heading(component_name, level=8)
                    self.add_table(rows, has_header=False)

        return self.save()

    def save(self):
        handle = BytesIO()
        self.document.save(handle)
        handle.seek(0)
        return handle.read()

    def get_row(self, field):
        return [field.name, field.value, '']

    def add_table(self, rows, has_header=False, **kwargs):
        table = self.document.add_table(rows=0, cols=0)
        for width in self.column_layout:
            table.add_column(width)

        if has_header:
            self.populate_row(table, rows.pop(0))

        for row in rows:
            self.populate_row(table, row, **kwargs)

        table.style = self.table_style
        return table

    def populate_row(self, table, row, **kwargs):
        cells = table.add_row().cells

        for i, (value, style) in enumerate(zip(row, self.row_sytle)):
            self.apply_style(cells[i], str(value), style)

    def apply_style(self, obj, value, style):
        run = obj.paragraphs[0].add_run(value)
        if not style:
            return run
        if style.font:
            run.font = style.font
        if style.size:
            run.font.size = style.size
        if style.bold:
            run.bold = style.bold
        if style.italic:
            run.italic = style.italic
        return run
